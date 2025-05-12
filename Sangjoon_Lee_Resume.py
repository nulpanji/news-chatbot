from docx import Document

doc = Document()
doc.add_heading('SANGJOON LEE', 0)
doc.add_paragraph('68, Mapo-daero, Mapo-gu, Seoul, South Korea')
doc.add_paragraph('+82 10 8735 7915 | nulpanji@gmail.com')
doc.add_paragraph()
doc.add_heading('PROFESSIONAL SUMMARY', 1)
doc.add_paragraph('Dynamic automotive and service industry leader with 20+ years of experience in after-sales, technical support, and business management. Proven expertise in operations, customer satisfaction, new business development, and digital transformation. Adept at leading cross-functional teams, optimizing processes, and driving organizational growth in both multinational corporations and start-up environments.')
doc.add_heading('CORE COMPETENCIES', 1)
doc.add_paragraph('• Automotive After-Sales & Service Operations\n• P&L & KPI Management\n• Team Leadership & Talent Development\n• Customer Experience & Complaint Resolution\n• Project & Event Management\n• Digital Transformation & System Integration\n• Cross-Cultural Communication')
doc.add_heading('PROFESSIONAL EXPERIENCE', 1)

# Sangsangwoori Co., Ltd. — Social Enterprise
doc.add_heading('Sangsangwoori Co., Ltd. — Social Enterprise', 2)
doc.add_paragraph('Career Division Lead / Social Innovation (Oct 2024 – Present)\n- Led ESG strategy projects for Hana Financial Group (“Power On Challenge – Second Life”), planning and executing programs for career transition and social impact.\n- Managed booth operations at the 2024 Gyeonggi Province Job Fair, coordinating logistics, staff, and outreach to maximize engagement.\n- Executed projects for the Korea Arts & Culture Committee, collaborating with government agencies to promote youth cultural spaces.')

doc.add_heading('Autopedia Co., Ltd. — IT & Automotive Start-up', 2)
doc.add_paragraph('Project Lead (Oct 2023 – May 2024)\n- Directed China EV import project: defined vehicle specs for the Korean market, negotiated pricing and warranty terms, and ensured regulatory compliance.\n- Managed service center operations in Seoul, Cheongju, and Seosan, optimizing cost structures and service quality.\n- Developed and implemented an integrated service center management system using Notion for workflow standardization and performance monitoring.\n- Conducted market analysis and competitor benchmarking to secure price competitiveness and superior warranty policies.')

doc.add_heading('Ojin Yanghaeng Co., Ltd. — Starbucks Official Service Provider', 2)
doc.add_paragraph('Head of Service Division (Jan 2022 – Aug 2022)\n- Managed service contracts for major clients including Starbucks, Ediya Coffee, McDonald’s, and Pizza Hut.\n- Introduced KPI-driven performance management, improving service efficiency and customer satisfaction.\n- Oversaw relocation and remodeling of 8 nationwide service centers, enhancing work environments and service quality.\n- Built an integrated management system to strengthen inter-branch collaboration and centralized control.')

doc.add_heading('MotorOne Co., Ltd. — Mercedes-Benz Official Dealer', 2)
doc.add_paragraph('General Manager, Service Division (Apr 2019 – Dec 2021)\n- Oversaw operations for 6 service centers in Northwest Gyeonggi, achieving record sales and expanding the service network.\n- Designed and implemented employee incentive and retention programs, resulting in increased staff stability and motivation.\n- Managed full P&L responsibility for the service division, optimizing resource allocation and operational efficiency.')

doc.add_heading('Hankook Tire & Technology Co., Ltd.', 2)
doc.add_paragraph('New Business Development Manager (Sep 2016 – Apr 2019)\n- Developed and launched JAX Motors, an imported car maintenance business, including site selection, facility setup, and staff recruitment.\n- Opened and managed 7 new branches, overseeing all aspects of P&L, facility management, and personnel.\n- Negotiated and executed dealer agreements with Peugeot & Citroen, acquiring and launching new service centers.')

doc.add_heading('AJ Networks Co., Ltd. — Jaguar & Land Rover Official Dealer', 2)
doc.add_paragraph('Service Branch Manager (Aug 2013 – Aug 2016)\n- Managed sales, profits, staff, and facilities for multiple service centers.\n- Set up 3 new service centers from the ground up, including equipment procurement and team training.\n- Awarded Best Service Manager in 2014 for outstanding operational performance and customer satisfaction.')

doc.add_heading('FMK Co., Ltd. — Ferrari & Maserati Official Importer', 2)
doc.add_paragraph('Service Planning & Branch Manager (Nov 2010 – Aug 2013)\n- Attracted new customers by restructuring parts pricing and introducing mobile service and exclusive Ferrari tow trucks.\n- Managed new car certification through both outsourcing and in-house processes.\n- Enhanced work efficiency by implementing on-site status boards and participating in technician training in Italy and Hong Kong.\n- Acted as a technical liaison with the headquarters in Italy.')

doc.add_heading('BMW Korea', 2)
doc.add_paragraph('After-Sales, Technical Support, Sales & Marketing (May 2001 – Sep 2010)\n- Coordinated government affairs, prepared recall campaign reports, and ensured regulatory compliance with the Ministry of Land, Infrastructure and Transport.\n- Handled customer complaints and legal cases, collaborating with law firms and consumer protection agencies.\n- Led the localization of user manuals and infotainment systems, managing translation vendors and working with BMW AG in Germany.\n- Served as technical hotline contact for nationwide dealers, providing solutions for critical vehicle issues (e.g., airbags, fire, sudden acceleration).\n- Participated in Technical Support and Customer Relations Committees (KAIDA), and managed escalated customer complaints and compensation cases.\n- Conducted new vehicle technical and environmental certifications, managed dealer marketing activities, and organized major brand events including Seoul/Busan Motor Shows and BMW Golf Cup.')

doc.add_heading('EDUCATION', 1)
doc.add_paragraph('• M.S. Industrial Engineering, Pittsburg State University, USA (1999)\n• B.S. Automotive Engineering, Pittsburg State University, USA (1998)\n• Automotive Maintenance Program, Apex Vocational School, NY, USA (1994)\n• Santa Monica College, USA (1992–1993)\n• Kyunggi High School, Korea (1989)')

doc.add_heading('DIGITAL SKILLS', 1)
doc.add_paragraph('• Collaboration: Slack, Jandi, Teams\n• Workspace: Google Drive, Notion\n• Mobile Office: Google Sheets, Slides, Notion, OneNote\n• Productivity: Apps Script & Python (AI), MS Office, Adobe Photoshop/Illustrator')

doc.save('Sangjoon_Lee_Resume.docx')
